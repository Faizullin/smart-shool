# from results.models import Result
# def save_student_answer_result_score(student_answer: StudentAnswer):
#     quiz = student_answer.question.quiz
#     exam = quiz.exam
#     student = student_answer.student
#     result = Result.objects.get(student=student, exam=exam)
#     quiz_questions = quiz.questions.all()
#     total_questions_count = quiz.questions_count
#     total_score_count = 0
#     user_answers = StudentAnswer.objects.filter(
#         question__in=quiz_questions, student=student)
#     print("OLD", result.theory_marks, ' for ', student_answer.score, ' + ',
#           [user_answer for user_answer in user_answers if user_answer.pk != student_answer.pk])
#     total_score_count = sum([(0 if user_answer.score is None else user_answer.score)
#                             for user_answer in user_answers if user_answer.pk != student_answer.pk])
#     total_score_count += student_answer.score
#     result.theory_marks = (total_score_count / total_questions_count) * 100
#     print("New", result.theory_marks)
#     result.save()
from pyapa import pyapa
from docx import Document


class FileChecker:
    def __init__(self):
        self.file = None
        self.content = None
        self.errors = {}
        self.init_errors()

    def init_errors(self):
        self.errors = {
            'abstract': [],
            'introduction': [],
            'discussion': [],
            'conclusion': [],
            'reference_list': [],
        }

    def validate_apa_style(self, content):
        a = pyapa.ApaCheck()
        text = u'Papaya are delicious fruit, it was concluded (Author et al. 2017).'
        match = a.match(text)
        # print(match[0].print())
        print("APA check:", [i for i in match])
        return False, match

    def validate_doc_title(self, content, key, label):
        for paragraph in doc.paragraphs:
            # Extract the content for different sections
            print("CJECK aprapghraph", paragraph,paragraph.text)
            # if label in paragraph.text:
            #     print("Found", key, label)
            # else:
            #     self.errors[key].append({
            #         "type": "NotFound",
            #         "message": "Not found"
            #     })
        return False, None

    def set_docx_file(self, file):
        self.content = file

    def validate_docx_file(self):
        steps = [
            {
                'method': self.validate_apa_style,
                'args': [self.content]
            },
            {
                'method': self.validate_doc_title,
                'args': [self.content, 'introduction', 'Introduction']
            },
            {
                'method': self.validate_doc_title,
                'args': [self.content, 'conclusion', 'Conclusion']
            },
            {
                'method': self.validate_doc_title,
                'args': [self.content, 'reference_list', 'Reference List']
            }
        ]
        for step in steps:
            is_correct, errors_list = step['method'](*step['args'])
        # found_error = False

        # for key in errors.keys():
        #     is_title, errors_list = self.validate_doc_title(content, key)
        #     if not is_title:
        #         errors[key] = errors_list
        #         found_error = True
        # is_apa_style_correct = self.validate_apa_style(content)
        # if not is_apa_style_correct:
        #     found_error = True
        #     errors['refrence_list'].append("Reference list is not correct")
        # return found_error, errors

doc = Document("test-data.docx")

# Initialize errors dictionary
errors = {
    "Introduction": [],
    "Abstract": [],
    "Reference List": []
}

# Function to check font size
def check_font_size(paragraphs, font_size):
    print([[paragraph.text, paragraph.style.name == 'Heading 1', paragraph.style.name, vars(paragraph.style)] for paragraph in paragraphs])
    return all(any(run.font.size != font_size for run in paragraph.runs) for paragraph in paragraphs)

# Function to count occurrences of a string
def count_occurrences(text: str, target):
    return text.count(target)

sections = {
    "Introduction": "Introduction",
    "Abstract": "Abstract",
    "Reference List": "Reference List"
}

for paragraph in doc.paragraphs:
    for section_name, section_text in sections.items():
        if section_text in paragraph.text:
            errors[section_name].append(f"Existence of {section_text} section: Passed")
            if not check_font_size(doc.paragraphs, 20):
                errors[section_name].append(f"{section_text} section font size is not 20px")
            if count_occurrences(paragraph.text, section_text) > 1:
                errors[section_name].append(f"Multiple occurrences of {section_text} section")

# Display errors
for section, error_list in errors.items():
    if error_list:
        print(f"Errors in {section}:")
        for error in error_list:
            print(f"- {error}")